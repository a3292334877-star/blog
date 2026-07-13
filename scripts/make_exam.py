#!/usr/bin/env python3
"""生成专插本计算机基础模拟卷 DOCX。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import argparse
from pathlib import Path


def main(output_path: Path):
    doc = Document()

    # -- 全局默认字体 --
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 页边距
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)


    def add_title(text, size=16):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        return p

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.font.size = Pt(11)
        return p


    def add_hint(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
        return p

    # ============================================================
    # 卷头
    # ============================================================
    add_title('广东省2026年普通专升本（专插本）招生考试')
    add_title('《计算机基础与程序设计》模拟卷（一）')
    add_subtitle('（考试时间：150分钟    满分：200分）')
    add_subtitle('命题依据：2024年真题回忆版 | 谭浩强《C语言程序设计》第5版 + 严蔚民《数据结构（C语言版）》')

    doc.add_paragraph()  # 空行

    # ============================================================
    # 一、单选题
    # ============================================================
    add_section_heading('一、单项选择题（本大题共20小题，每小题3分，共60分）')
    add_body('在每小题列出的四个备选项中只有一个是符合题目要求的，请将其代码填写在题后的括号内。错选、多选或未选均无分。')

    danxuan = [
        ('1. 一个C程序的执行是从（ ）。',
         ['A. 本程序的main函数开始，到main函数结束',
          'B. 本程序的第一个函数开始，到最后一个函数结束',
          'C. 本程序的main函数开始，到最后一个函数结束',
          'D. 本程序的第一个函数开始，到main函数结束']),

        ('2. 以下选项中，合法的C语言标识符是（ ）。',
         ['A. 3day', 'B. day-3', 'C. _day3', 'D. int']),

        ('3. 若有定义：int a=5, b=3; 则表达式 a % b + b / a 的值是（ ）。',
         ['A. 2', 'B. 3', 'C. 1', 'D. 2.6']),

        ('4. 以下能正确表示数学式 0 ≤ x < 20 的C语言表达式是（ ）。',
         ['A. 0<=x<20', 'B. x>=0||x<20', 'C. x>=0 && x<20', 'D. 0<=x && <20']),

        ('5. 若变量c为char类型，能正确判断c为小写字母的表达式是（ ）。',
         ['A. \'a\'<=c<=\'z\'', 'B. c>=\'a\' && c<=\'z\'',
          'C. \'a\'<=c || c<=\'z\'', 'D. c>=\'a\' || c<=\'z\'']),

        ('6. 设有定义 int a[10]={1,2,3,4,5,6,7,8,9,10}, *p=&a[3]; 则 p[2] 的值是（ ）。',
         ['A. 3', 'B. 4', 'C. 5', 'D. 6']),

        ('7. 以下程序的输出结果是（ ）。\n    int a=5, b=6;\n    printf("%d\\n", (a++, ++b, a+b));',
         ['A. 11', 'B. 12', 'C. 13', 'D. 14']),

        ('8. 以下程序段的输出结果是（ ）。\n    int x=3;\n    do { printf("%d", x--); } while (!x);',
         ['A. 3', 'B. 321', 'C. 32', 'D. 3然后死循环']),

        ('9. 若有定义 int b[3][4]={0}; 以下叙述正确的是（ ）。',
         ['A. 只有元素 b[0][0] 可得到初值0',
          'B. 此定义语句不正确',
          'C. 数组b中各元素都可得到初值0，但无法确定是多少',
          'D. 数组b中每个元素均可得到初值0']),

        ('10. 以下叙述中正确的是（ ）。',
         ['A. break语句只能用于switch语句',
          'B. continue语句只能用于循环语句',
          'C. break语句只能用于循环语句',
          'D. 在switch语句中必须使用default']),

        ('11. 设有定义 char str[]="Hello"; 则 sizeof(str) 和 strlen(str) 的值分别是（ ）。',
         ['A. 5和5', 'B. 5和6', 'C. 6和5', 'D. 6和6']),

        ('12. 以下函数的功能是（ ）。\n    int fun(char *s) {\n        char *t=s;\n        while(*t) t++;\n        return t-s;\n    }',
         ['A. 比较两个字符串的大小', 'B. 计算字符串的长度',
          'C. 将字符串s复制到t', 'D. 连接两个字符串']),

        ('13. C语言中结构体类型变量在程序执行期间（ ）。',
         ['A. 所有成员一直驻留在内存中', 'B. 只有一个成员驻留在内存中',
          'C. 部分成员驻留在内存中', 'D. 没有成员驻留在内存中']),

        ('14. 在单链表中，指针p指向某结点，则删除该结点后继结点的正确操作是（ ）。',
         ['A. p->next = p->next->next',
          'B. p = p->next; p->next = p->next->next',
          'C. p->next = p->next',
          'D. p = p->next->next']),

        ('15. 栈和队列的共同点是（ ）。',
         ['A. 都是先进后出', 'B. 都是先进先出',
          'C. 只允许在端点处插入和删除元素', 'D. 没有共同点']),

        ('16. 一个栈的入栈序列为1,2,3,4,5，则以下不可能的出栈序列是（ ）。',
         ['A. 2,3,4,1,5', 'B. 5,4,3,2,1', 'C. 3,1,4,2,5', 'D. 1,2,3,4,5']),

        ('17. 在一个具有n个顶点的无向完全图中，边的总数为（ ）。',
         ['A. n(n-1)', 'B. n(n-1)/2', 'C. n²', 'D. n-1']),

        ('18. 对长度为n的线性表进行顺序查找，在等概率情况下，查找成功的平均查找长度为（ ）。',
         ['A. n', 'B. n/2', 'C. (n+1)/2', 'D. (n-1)/2']),

        ('19. 深度为4的二叉树最多有（ ）个结点。',
         ['A. 8', 'B. 15', 'C. 16', 'D. 7']),

        ('20. 数据的逻辑结构可分为（ ）。',
         ['A. 动态结构和静态结构', 'B. 紧凑结构和非紧凑结构',
          'C. 线性结构和非线性结构', 'D. 内部结构和外部结构']),
    ]

    for title, options in danxuan:
        add_body(title)
        for opt in options:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.line_spacing = 1.4
            run = p.add_run(opt)
            run.font.size = Pt(11)
        doc.add_paragraph()  # 空行

    # ============================================================
    # 二、判断题
    # ============================================================
    add_section_heading('二、判断题（本大题共10小题，每小题2分，共20分）')
    add_body('判断下列各题正误，正确的在题后括号内打"√"，错误的打"×"。（不需要写解析，只标记对错即可）')

    panduan = [
        '1. 在C语言中，注释可以嵌套使用。（ ）',
        '2. 使用 for 循环时可以省略循环体，此时循环没有实际意义。（ ）',
        '3. C语言中，函数可以嵌套定义。（ ）',
        '4. 一个数组中的所有元素在内存中是连续存储的。（ ）',
        '5. 静态局部变量在函数调用结束后仍然保留其值。（ ）',
        '6. 指针变量中存放的是所指向变量的值。（ ）',
        '7. 链式存储结构比顺序存储结构更节省存储空间。（ ）',
        '8. 串是一种特殊的线性表，其数据元素只能是字符。（ ）',
        '9. 二叉树的第i层上最多有 2^(i-1) 个结点（i≥1）。（ ）',
        '10. 算法的五个特性包括：有穷性、确定性、可行性、有零个或多个输入、有零个或多个输出。（ ）',
    ]

    for t in panduan:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        p.add_run(t)

    # ============================================================
    # 三、填空题
    # ============================================================
    add_section_heading('三、填空题（本大题共5小题，每小题4分，共20分）')

    tiankong = [
        '1. 表达式  5 + 2 * 7 % 3  的值是 ________。',
        '2. 下面程序段执行后，变量 s 的值是 ________。\n    int a=10, b=20, s=0;\n    if (a<b) s=a+b;\n    else s=a-b;',
        '3. C语言中，动态内存分配函数 malloc 的返回值类型是 ________。',
        '4. n个顶点的连通图至少有 ________ 条边。',
        '5. 长度为n的顺序表中，删除第i个元素需要向前移动 ________ 个元素。',
    ]

    for t in tiankong:
        add_body(t)

    # ============================================================
    # 四、简答题
    # ============================================================
    add_section_heading('四、简答题（本大题共4小题，共40分）')

    jianda = [
        ('1.（10分）简述C语言中 break 和 continue 的区别，并说明它们各自的使用场景。',
         ''),
        ('2.（10分）什么是顺序存储结构？什么是链式存储结构？请从空间利用率、查找和插入删除三个方面对比两者的优缺点。',
         ''),
        ('3.（10分）阅读以下程序段，回答问题：\n\n    int a[5] = {10, 20, 30, 40, 50};\n    int *p = a;\n    printf("%d\\n", *(p+2));\n    printf("%d\\n", *p+2);\n\n（1）第一条 printf 输出什么？（5分）\n（2）第二条 printf 输出什么？（5分）\n（3）*(p+2) 和 *p+2 有什么区别？（附加说明）',
         ''),
        ('4.（10分）简述栈和队列的定义，说明它们各自属于哪种操作受限的线性表，并各举一个实际应用场景。',
         ''),
    ]

    for title, _ in jianda:
        add_body(title)

    # ============================================================
    # 五、计算/程序分析题
    # ============================================================
    add_section_heading('五、计算/程序分析题（本大题共3小题，每小题10分，共30分）')
    add_hint('提示：写出运行结果即可，不需要写代码。')

    jisuan = [
        ('1.（10分）写出以下程序的运行结果：\n\n#include <stdio.h>\n\nint fun(int n) {\n    if (n <= 1) return 1;\n    return n * fun(n - 1);\n}\n\nint main() {\n    printf("%d\\n", fun(5));\n    return 0;\n}',
         ''),
        ('2.（10分）写出以下程序的运行结果：\n\n#include <stdio.h>\n\nint main() {\n    int i, j;\n    for (i = 1; i <= 4; i++) {\n        for (j = 1; j <= i; j++) {\n            printf("%d", j);\n        }\n        printf("\\n");\n    }\n    return 0;\n}',
         ''),
        ('3.（10分）写出以下程序的运行结果：\n\n#include <stdio.h>\n#include <string.h>\n\nint main() {\n    char s[] = "a1b2c3d4";\n    int sum = 0;\n    for (int i = 0; s[i] != \'\\0\'; i++) {\n        if (s[i] >= \'0\' && s[i] <= \'9\') {\n            sum += s[i] - \'0\';\n        }\n    }\n    printf("%d\\n", sum);\n    return 0;\n}',
         ''),
    ]

    for title, _ in jisuan:
        add_body(title)

    # ============================================================
    # 六、应用题（编程题）
    # ============================================================
    add_section_heading('六、应用题（本大题共3小题，每小题10分，共30分）')
    add_hint('提示：请编写完整的C语言程序，包含必要的头文件和main函数。')

    yingyong = [
        ('1.（10分）编写程序，输出 1 到 100 之间所有能被 3 整除但不能被 5 整除的数，每行输出 5 个数。',
         ''),
        ('2.（10分）编写一个函数 int max(int a[], int n)，返回数组 a 中 n 个元素的最大值。并在main函数中定义一个包含10个元素的数组，从键盘输入10个整数，调用max函数求出最大值并输出。',
         ''),
        ('3.（10分）定义一个结构体类型 Book，包含书名（char name[50]）、作者（char author[30]）、价格（float price）三个成员。从键盘输入3本书的信息，找出价格最低的书，输出其所有信息。',
         ''),
    ]

    for title, _ in yingyong:
        add_body(title)

    # ============================================================
    # 保存
    # ============================================================
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ 已保存到: {output_path}')

def parse_args():
    default_output = Path.home() / 'Desktop' / '专插本计算机基础模拟卷（一）.docx'
    parser = argparse.ArgumentParser(description='生成专插本计算机基础模拟卷 DOCX')
    parser.add_argument(
        '--output',
        type=Path,
        default=default_output,
        help=f'输出文件路径（默认：{default_output}）',
    )
    return parser.parse_args()


if __name__ == '__main__':
    args = parse_args()
    main(args.output.expanduser().resolve())